# # from __future__ import annotations

# # import json
# # from pathlib import Path
# # from typing import Any


# # def read_text(path: str | Path) -> str:
# #     p = Path(path)
# #     return p.read_text(encoding="utf-8")


# # def write_json(path: str | Path, data: Any) -> None:
# #     p = Path(path)
# #     p.parent.mkdir(parents=True, exist_ok=True)
# #     p.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


# # def read_json(path: str | Path, default: Any) -> Any:
# #     p = Path(path)
# #     if not p.exists():
# #         return default
# #     return json.loads(p.read_text(encoding="utf-8"))

# from __future__ import annotations

# import json
# from pathlib import Path
# from typing import Any


# def read_text(path: str | Path) -> str:
#     p = Path(path)
#     suffix = p.suffix.lower()

#     if suffix in {".txt", ".md", ".csv", ".json", ".yaml", ".yml"}:
#         return p.read_text(encoding="utf-8", errors="ignore")

#     if suffix == ".pdf":
#         return _read_pdf_text(p)

#     if suffix == ".docx":
#         return _read_docx_text(p)

#     # fallback
#     return p.read_text(encoding="utf-8", errors="ignore")


# def _read_pdf_text(path: Path) -> str:
#     text_parts: list[str] = []

#     # First try pdfplumber for better layout extraction
#     try:
#         import pdfplumber  # type: ignore

#         with pdfplumber.open(path) as pdf:
#             for page in pdf.pages:
#                 page_text = page.extract_text() or ""
#                 if page_text.strip():
#                     text_parts.append(page_text)
#     except Exception:
#         pass

#     if text_parts:
#         return "\n\n".join(text_parts).strip()

#     # Fallback to pypdf
#     try:
#         from pypdf import PdfReader  # type: ignore

#         reader = PdfReader(str(path))
#         for page in reader.pages:
#             page_text = page.extract_text() or ""
#             if page_text.strip():
#                 text_parts.append(page_text)
#     except Exception:
#         pass

#     if text_parts:
#         return "\n\n".join(text_parts).strip()

#     raise ValueError(f"Could not extract text from PDF: {path}")


# def _read_docx_text(path: Path) -> str:
#     try:
#         from docx import Document  # type: ignore

#         doc = Document(str(path))
#         lines = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
#         return "\n".join(lines).strip()
#     except Exception as exc:
#         raise ValueError(f"Could not extract text from DOCX: {path}") from exc


# def write_json(path: str | Path, data: Any) -> None:
#     p = Path(path)
#     p.parent.mkdir(parents=True, exist_ok=True)
#     p.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


# def read_json(path: str | Path, default: Any) -> Any:
#     p = Path(path)
#     if not p.exists():
#         return default
#     return json.loads(p.read_text(encoding="utf-8"))

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import yaml


def read_text(path: str | Path) -> str:
    p = Path(path)
    suffix = p.suffix.lower()

    if suffix in {".txt", ".md", ".csv", ".json", ".yaml", ".yml"}:
        return p.read_text(encoding="utf-8", errors="ignore")

    if suffix == ".pdf":
        return _read_pdf_text(p)

    if suffix == ".docx":
        return _read_docx_text(p)

    return p.read_text(encoding="utf-8", errors="ignore")


def _read_pdf_text(path: Path) -> str:
    text_parts: list[str] = []

    try:
        import pdfplumber  # type: ignore

        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(page_text)
    except Exception:
        pass

    if text_parts:
        return "\n\n".join(text_parts).strip()

    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(page_text)
    except Exception:
        pass

    if text_parts:
        return "\n\n".join(text_parts).strip()

    raise ValueError(f"Could not extract text from PDF: {path}")


def _read_docx_text(path: Path) -> str:
    try:
        from docx import Document  # type: ignore

        doc = Document(str(path))
        lines = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        return "\n".join(lines).strip()
    except Exception as exc:
        raise ValueError(f"Could not extract text from DOCX: {path}") from exc


def write_json(path: str | Path, data: Any) -> None:
    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


def read_json(path: str | Path, default: Any) -> Any:
    p = Path(path)
    if not p.exists():
        return default
    return json.loads(p.read_text(encoding="utf-8"))


def read_yaml(path: str | Path, default: Any | None = None) -> Any:
    p = Path(path)
    if not p.exists():
        if default is not None:
            return default
        raise FileNotFoundError(f"YAML file not found: {path}")

    with p.open("r", encoding="utf-8") as f:
        data = yaml.safe_load(f)

    if data is None:
        return default if default is not None else {}

    return data